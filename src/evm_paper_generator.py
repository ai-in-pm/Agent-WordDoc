from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os
from dotenv import load_dotenv
import openai
import time

class EVMPaperGenerator:
    def __init__(self, topic="Earned Value Management", template=None):
        self.document = Document(template) if template else Document()
        self.topic = topic
        self.setup_document()
        self.setup_openai()
        self.citations = []
        self.references = []
        self.progress = {
            'status': 'initializing',
            'current_section': None,
            'total_sections': 0,
            'completed_sections': 0
        }
        
        # Default sections that can be overridden
        self.sections = [
            "Abstract",
            "Introduction",
            "Literature Review",
            "Key Concepts",
            "Methodology",
            "Results and Discussion",
            "Conclusion",
            "References"
        ]
        
        # Define default tables
        self.tables = [
            {"section": "Key Concepts", "title": f"Key {topic} Metrics"},
            {"section": "Methodology", "title": "Data Collection Methods"},
            {"section": "Results and Discussion", "title": "Comparative Analysis"}
        ]
        
        # Define default figures
        self.figures = [
            {"section": "Key Concepts", "title": f"{topic} Overview Diagram"},
            {"section": "Results and Discussion", "title": "Performance Analysis"},
            {"section": "Results and Discussion", "title": "Trend Analysis"}
        ]

    def handle_error(self, operation, error):
        """Handle errors with retry logic and logging"""
        print(f"[ERROR] {operation} failed: {str(error)}")
        
        retries = 0
        max_retries = 3
        while retries < max_retries:
            try:
                # Attempt to recover
                if operation == "API Call":
                    return self.generate_section_content(self.progress['current_section'])
                elif operation == "Document Write":
                    return self.write_to_document(self.progress['current_section'])
                
                return True
                
            except Exception as recovery_error:
                retries += 1
                print(f"[RETRY] Attempt {retries}/{max_retries} for {operation}")
                time.sleep(2)
        
        print(f"[FAILED] {operation} after {max_retries} attempts")
        return False

    def update_progress(self, section_name=None, status=None):
        """Update and report progress"""
        if section_name:
            self.progress['current_section'] = section_name
        if status:
            self.progress['status'] = status

        print(f"\n=== Progress Update ===")
        print(f"Status: {self.progress['status']}")
        print(f"Current Section: {self.progress['current_section']}")
        print(f"Completed: {self.progress['completed_sections']}/{self.progress['total_sections']}")
        print("======================\n")

    def validate_content(self, content):
        """Validate generated content"""
        if not content or content.strip() == "":
            raise ValueError("Generated content is empty")

        # Check for minimum content length
        if len(content.split()) < 50:
            raise ValueError("Generated content is too short")

        # Check for proper formatting
        if not any(c.isupper() for c in content):
            raise ValueError("Generated content lacks proper capitalization")

        return True

    def generate_section_content(self, section_name):
        """Generate content for a section with error handling"""
        try:
            self.update_progress(section_name, "generating")
            
            prompt = f"Write a detailed academic section on {self.topic} for the '{section_name}' section. Focus on professional academic tone, include relevant citations in APA format, and ensure thorough coverage of the topic."
            
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": f"You are an expert academic writer specializing in {self.topic}. Write in a formal academic style with proper APA citations. When citing sources, use the format (Author, Year) and add the full reference to your response prefixed with [REFERENCE]."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            
            content = response.choices[0].message.content
            
            # Validate content
            self.validate_content(content)
            
            # Extract references
            for line in content.split('\n'):
                if line.startswith('[REFERENCE]'):
                    ref = line.replace('[REFERENCE]', '').strip()
                    if ref not in self.references:
                        self.references.append(ref)
            
            # Remove reference lines from content
            content = '\n'.join(line for line in content.split('\n') if not line.startswith('[REFERENCE]'))
            
            self.update_progress(section_name, "completed")
            return content
            
        except Exception as e:
            self.handle_error("API Call", e)
            raise

    def write_to_document(self, section_name, content):
        """Write content to document with error handling"""
        try:
            self.update_progress(section_name, "writing")
            
            # Add section heading
            heading = self.document.add_heading(section_name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content
            paragraph = self.document.add_paragraph(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add spacing
            self.document.add_paragraph()
            
            self.progress['completed_sections'] += 1
            self.update_progress(section_name, "written")
            return True
            
        except Exception as e:
            self.handle_error("Document Write", e)
            raise

    def generate_paper(self):
        """Generate the complete paper"""
        try:
            self.progress['total_sections'] = len(self.sections)
            self.progress['completed_sections'] = 0
            
            for section in self.sections:
                if section == "References":
                    # Handle references separately
                    self.write_references()
                    continue
                
                # Generate content
                content = self.generate_section_content(section)
                
                # Write to document
                self.write_to_document(section, content)
                
            # Save the document
            self.save_document()
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Failed to generate paper: {str(e)}")
            raise

    def write_references(self):
        """Write the references section"""
        try:
            # Add references heading
            heading = self.document.add_heading("References", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add references
            for ref in sorted(self.references):
                paragraph = self.document.add_paragraph(ref)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            # Add spacing
            self.document.add_paragraph()
            
            self.progress['completed_sections'] += 1
            self.update_progress("References", "written")
            return True
            
        except Exception as e:
            self.handle_error("References Write", e)
            raise

    def save_document(self):
        """Save the document with error handling"""
        try:
            # Create output directory if it doesn't exist
            output_dir = "output"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename with timestamp
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{self.topic.replace(' ', '_')}_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            # Save document
            self.document.save(filepath)
            print(f"\n=== Document Saved ===")
            print(f"File: {filename}")
            print(f"Location: {os.path.abspath(filepath)}")
            print("====================\n")
            
            return filepath
            
        except Exception as e:
            self.handle_error("Document Save", e)
            raise

    def setup_document(self):
        # Set up margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def setup_openai(self):
        load_dotenv()
        openai.api_key = os.getenv('OPENAI_API_KEY')

    def add_title_page(self):
        # Title
        title = self.document.add_paragraph()
        title_run = title.add_run(f"{self.topic}:\nA Comprehensive Analysis")
        title_run.font.size = Pt(24)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some space
        self.document.add_paragraph()
        
        # Author
        author = self.document.add_paragraph()
        author.add_run("Generated by AI Assistant")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date = self.document.add_paragraph()
        date.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_page_break()

def main():
    # Example usage with custom topic
    generator = EVMPaperGenerator(topic="Artificial Intelligence in Healthcare")
    generator.add_title_page()
    generator.generate_paper()
    print(f"Academic paper on {generator.topic} has been generated successfully!")

if __name__ == "__main__":
    main()
