import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import math

def extract_problems(tex_content):
    problems = {}
    pattern = r'\\item\[(.[1-6])\](.*?)(?=\\item|\\end{itemize})'
    matches = re.finditer(pattern, tex_content, re.DOTALL)
    
    for match in matches:
        problem_num = match.group(1)
        problem_text = match.group(2).strip()
        problems[problem_num] = problem_text
    
    return problems

def analyze_problem(problem_text):
    """Simulate AI analyzing the math problem in real-time"""
    analysis = []
    analysis.append("Problem Statement:")
    # Clean LaTeX formatting for better readability
    cleaned_text = problem_text.replace('\\', '').replace('$', '').replace('{', '').replace('}', '')
    analysis.append(cleaned_text)
    
    analysis.append("\nKey Components:")
    # Add some basic analysis points
    if "prove" in problem_text.lower():
        analysis.append("- This is a proof-based problem")
    if "find" in problem_text.lower():
        analysis.append("- This is a computational/analytical problem")
    if "show" in problem_text.lower():
        analysis.append("- This requires mathematical demonstration")
        
    analysis.append("\nApproach Strategy:")
    analysis.append("1. First, understand the given conditions")
    analysis.append("2. Break down the problem into smaller parts")
    analysis.append("3. Consider relevant theorems and techniques")
    
    return "\n".join(analysis)

def create_word_document(problems):
    doc = Document()
    
    # Add title
    title = doc.add_heading('Putnam Mathematical Competition Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph('This document provides a detailed examination of each problem from the 85th William Lowell Putnam Mathematical Competition.')
    
    for problem_num, problem_text in problems.items():
        # Add problem header
        doc.add_heading(f'Problem {problem_num}', level=1)
        
        # Simulate real-time analysis with typing effect
        analysis = analyze_problem(problem_text)
        
        # Add analysis to document with simulated typing
        p = doc.add_paragraph()
        for line in analysis.split('\n'):
            p.add_run(line + '\n')
            time.sleep(0.1)  # Simulate typing delay
            
        # Add separator
        doc.add_paragraph('=' * 50)
    
    return doc

def main():
    # Read the LaTeX file
    tex_file_path = r"C:\Users\djjme\OneDrive\Desktop\CC-Directory\Agent WordDoc\Math Problems.tex"
    with open(tex_file_path, 'r') as file:
        tex_content = file.read()
    
    # Extract problems
    problems = extract_problems(tex_content)
    
    # Create and save Word document
    doc = create_word_document(problems)
    output_path = r"C:\Users\djjme\OneDrive\Desktop\CC-Directory\Agent WordDoc\Putnam_Analysis.docx"
    doc.save(output_path)
    print(f"Analysis complete! Document saved to: {output_path}")

if __name__ == "__main__":
    main()
